# core/file_processors.py
# Datapizza v1.5.0 - Processori per file uploadati in chat
# ============================================================================
#
# Questo modulo gestisce l'estrazione di testo da documenti e la
# preparazione di immagini per modelli Vision.
#
# Tipi supportati:
# - Documenti: PDF, TXT, MD, DOCX
# - Immagini: PNG, JPG, JPEG, GIF, WEBP (solo modelli Vision)
#
# ============================================================================

import io
import base64
from pathlib import Path
from typing import Optional, Tuple, List
from dataclasses import dataclass, field

from config import MAX_DOCUMENT_CHARS


@dataclass
class ProcessedFile:
    """
    Risultato del processamento di un file uploadato.
    
    Attributes:
        filename: Nome originale del file
        file_type: "document" | "image" | "unknown"
        content: Testo estratto (documenti) o base64 (immagini)
        mime_type: MIME type del file
        size_bytes: Dimensione in bytes
        preview: Anteprima per UI (primi N caratteri o thumbnail base64)
        error: Messaggio di errore (se presente)
    """
    filename: str
    file_type: str
    content: str
    mime_type: str
    size_bytes: int
    preview: str = ""
    error: Optional[str] = None


# ============================================================================
# UTILITY FUNCTIONS
# ============================================================================

def is_image_file(filename: str) -> bool:
    """
    Verifica se il file è un'immagine supportata.
    
    Args:
        filename: Nome del file
        
    Returns:
        True se è un'immagine supportata
    """
    ext = Path(filename).suffix.lower()
    return ext in [".png", ".jpg", ".jpeg", ".gif", ".webp"]


def is_document_file(filename: str) -> bool:
    """
    Verifica se il file è un documento supportato.
    
    Args:
        filename: Nome del file
        
    Returns:
        True se è un documento supportato
    """
    ext = Path(filename).suffix.lower()
    return ext in [".pdf", ".txt", ".md", ".docx"]


def get_file_extension(filename: str) -> str:
    """Estrae l'estensione in lowercase."""
    return Path(filename).suffix.lower()


# ============================================================================
# DOCUMENT EXTRACTORS
# ============================================================================

def extract_text_from_txt(file_bytes: bytes, encoding: str = "utf-8") -> str:
    """
    Estrae testo da file TXT o MD.
    
    Prova UTF-8, poi Latin-1 come fallback.
    
    Args:
        file_bytes: Contenuto del file in bytes
        encoding: Encoding preferito
        
    Returns:
        Testo estratto
    """
    try:
        return file_bytes.decode(encoding)
    except UnicodeDecodeError:
        try:
            return file_bytes.decode("latin-1")
        except Exception:
            return file_bytes.decode("utf-8", errors="ignore")


def extract_text_from_pdf(file_bytes: bytes) -> str:
    """
    Estrae testo da PDF usando PyPDF2.
    
    PyPDF2 è già nelle dipendenze del progetto.
    
    Args:
        file_bytes: Contenuto del PDF in bytes
        
    Returns:
        Testo estratto con indicazione pagine
    """
    try:
        from PyPDF2 import PdfReader
        
        reader = PdfReader(io.BytesIO(file_bytes))
        text_parts = []
        
        for page_num, page in enumerate(reader.pages, 1):
            page_text = page.extract_text()
            if page_text and page_text.strip():
                text_parts.append(f"[Pagina {page_num}]\n{page_text.strip()}")
        
        if not text_parts:
            return "[PDF senza testo estraibile - potrebbe contenere solo immagini]"
        
        return "\n\n".join(text_parts)
    
    except ImportError:
        return "[Errore: PyPDF2 non installato. Installa con: pip install PyPDF2]"
    except Exception as e:
        return f"[Errore estrazione PDF: {e}]"


def extract_text_from_docx(file_bytes: bytes) -> str:
    """
    Estrae testo da DOCX usando python-docx.
    
    Estrae sia paragrafi che contenuto delle tabelle.
    
    Requisito: pip install python-docx
    
    Args:
        file_bytes: Contenuto del DOCX in bytes
        
    Returns:
        Testo estratto
    """
    try:
        from docx import Document
        
        doc = Document(io.BytesIO(file_bytes))
        paragraphs = []
        
        # Estrai paragrafi
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                paragraphs.append(text)
        
        # Estrai tabelle
        for table in doc.tables:
            table_rows = []
            for row in table.rows:
                row_cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if row_cells:
                    table_rows.append(" | ".join(row_cells))
            if table_rows:
                paragraphs.append("\n".join(table_rows))
        
        if not paragraphs:
            return "[Documento DOCX vuoto]"
        
        return "\n\n".join(paragraphs)
    
    except ImportError:
        return "[Errore: python-docx non installato. Installa con: pip install python-docx]"
    except Exception as e:
        return f"[Errore estrazione DOCX: {e}]"


# ============================================================================
# IMAGE PROCESSING
# ============================================================================

def process_image_to_base64(file_bytes: bytes, filename: str) -> Tuple[str, str]:
    """
    Converte immagine in base64 per modelli Vision.
    
    Args:
        file_bytes: Contenuto dell'immagine in bytes
        filename: Nome del file (per determinare MIME type)
        
    Returns:
        Tupla (base64_data, mime_type)
    """
    ext = get_file_extension(filename)
    
    mime_map = {
        ".png": "image/png",
        ".jpg": "image/jpeg",
        ".jpeg": "image/jpeg",
        ".gif": "image/gif",
        ".webp": "image/webp",
    }
    
    mime_type = mime_map.get(ext, "image/png")
    base64_data = base64.b64encode(file_bytes).decode("utf-8")
    
    return base64_data, mime_type


def create_image_thumbnail(file_bytes: bytes, max_size: int = 200) -> Optional[str]:
    """
    Crea thumbnail base64 per anteprima UI.
    
    Requisito: Pillow (opzionale - se non presente, ritorna None)
    
    Args:
        file_bytes: Contenuto dell'immagine in bytes
        max_size: Dimensione massima lato (default 200px)
        
    Returns:
        Thumbnail in base64 o None se Pillow non disponibile
    """
    try:
        from PIL import Image
        
        img = Image.open(io.BytesIO(file_bytes))
        
        # Converti RGBA in RGB se necessario (per JPEG)
        if img.mode in ('RGBA', 'LA', 'P'):
            background = Image.new('RGB', img.size, (255, 255, 255))
            if img.mode == 'P':
                img = img.convert('RGBA')
            background.paste(img, mask=img.split()[-1] if img.mode == 'RGBA' else None)
            img = background
        
        img.thumbnail((max_size, max_size), Image.Resampling.LANCZOS)
        
        buffer = io.BytesIO()
        img.save(buffer, format="PNG", optimize=True)
        buffer.seek(0)
        
        return base64.b64encode(buffer.getvalue()).decode("utf-8")
    
    except ImportError:
        # Pillow non installato - ritorna None
        return None
    except Exception:
        return None


# ============================================================================
# MAIN PROCESSOR
# ============================================================================

def process_uploaded_file(uploaded_file) -> ProcessedFile:
    """
    Processa un file uploadato da Streamlit.
    
    Determina automaticamente il tipo di file e applica
    l'estrattore appropriato.
    
    Args:
        uploaded_file: Oggetto UploadedFile da st.file_uploader
        
    Returns:
        ProcessedFile con contenuto estratto/processato
    """
    filename = uploaded_file.name
    file_bytes = uploaded_file.read()
    size_bytes = len(file_bytes)
    ext = get_file_extension(filename)
    
    # Reset buffer per successive letture
    uploaded_file.seek(0)
    
    # ========== IMMAGINI ==========
    if is_image_file(filename):
        base64_data, mime_type = process_image_to_base64(file_bytes, filename)
        thumbnail = create_image_thumbnail(file_bytes)
        
        return ProcessedFile(
            filename=filename,
            file_type="image",
            content=base64_data,
            mime_type=mime_type,
            size_bytes=size_bytes,
            preview=thumbnail if thumbnail else "",
            error=None
        )
    
    # ========== DOCUMENTI ==========
    if is_document_file(filename):
        # Estrai testo in base al tipo
        if ext == ".pdf":
            text = extract_text_from_pdf(file_bytes)
            mime_type = "application/pdf"
        elif ext == ".docx":
            text = extract_text_from_docx(file_bytes)
            mime_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        elif ext == ".md":
            text = extract_text_from_txt(file_bytes)
            mime_type = "text/markdown"
        else:  # .txt
            text = extract_text_from_txt(file_bytes)
            mime_type = "text/plain"
        
        # Verifica se c'è stato un errore nell'estrazione
        error = None
        if text.startswith("[Errore"):
            error = text
        
        # Tronca se troppo lungo
        original_len = len(text)
        if original_len > MAX_DOCUMENT_CHARS:
            text = text[:MAX_DOCUMENT_CHARS] + f"\n\n[... troncato a {MAX_DOCUMENT_CHARS:,} caratteri su {original_len:,} totali]"
        
        # Preview: primi 500 caratteri
        preview = text[:500] + "..." if len(text) > 500 else text
        
        return ProcessedFile(
            filename=filename,
            file_type="document",
            content=text,
            mime_type=mime_type,
            size_bytes=size_bytes,
            preview=preview,
            error=error
        )
    
    # ========== TIPO NON SUPPORTATO ==========
    return ProcessedFile(
        filename=filename,
        file_type="unknown",
        content="",
        mime_type="application/octet-stream",
        size_bytes=size_bytes,
        preview="",
        error=f"Tipo file non supportato: {ext}"
    )


def process_multiple_files(uploaded_files: List) -> List[ProcessedFile]:
    """
    Processa multipli file uploadati.
    
    Args:
        uploaded_files: Lista di UploadedFile da st.file_uploader
        
    Returns:
        Lista di ProcessedFile
    """
    processed = []
    for f in uploaded_files:
        if f is not None:
            processed.append(process_uploaded_file(f))
    return processed


# ============================================================================
# PROMPT BUILDING HELPERS
# ============================================================================

def build_document_context(processed_files: List[ProcessedFile]) -> str:
    """
    Costruisce il contesto testuale dai documenti per il prompt.
    
    Args:
        processed_files: Lista di file processati
        
    Returns:
        Stringa con contenuto documenti formattato
    """
    documents = [f for f in processed_files if f.file_type == "document" and not f.error]
    
    if not documents:
        return ""
    
    parts = []
    for doc in documents:
        parts.append(f"[📄 File: {doc.filename}]\n{doc.content}")
    
    return "\n\n--- FILE ALLEGATO ---\n".join(parts)


def get_images_for_vision(processed_files: List[ProcessedFile]) -> List[dict]:
    """
    Prepara le immagini nel formato per API Vision.
    
    Args:
        processed_files: Lista di file processati
        
    Returns:
        Lista di dict per API multimodal
    """
    images = [f for f in processed_files if f.file_type == "image" and not f.error]
    
    return [
        {
            "filename": img.filename,
            "base64": img.content,
            "mime_type": img.mime_type
        }
        for img in images
    ]


def get_attachment_names(processed_files: List[ProcessedFile]) -> List[str]:
    """
    Estrae i nomi dei file per salvare nei metadati del messaggio.
    
    Args:
        processed_files: Lista di file processati
        
    Returns:
        Lista di nomi file
    """
    return [f.filename for f in processed_files if not f.error]
